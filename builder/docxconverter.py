# -*- coding: utf-8 -*-
__author__ = 'Bárdos Dávid'

from sys import argv
from converter import Converter
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


class DocxConverter(Converter):
    def __init__(self, output_file_name):
        self._document = Document()
        self._raw_document_content = []
        self._output_file_path = "target/{}.docx".format(output_file_name)
        self._raw_metadata = self._read_metadata()
        self._grey = "BEBEBE"

    def convert(self):
        self._set_page_size_to_a4()
        self._generate_styles()
        sources = self._process_input_files()
        for source in sources:
            self._raw_document_content += source
        self._process_raw_document_content()
        self._document.save(self._output_file_path)

    def _set_page_size_to_a4(self):
        zeroth_section = self._document.sections[0]
        zeroth_section.page_height = Mm(297)
        zeroth_section.page_width = Mm(210)

    def _generate_styles(self):
        normal = self._document.styles["Normal"]
        normal.font.name = "Times New Roman"
        body = self._document.styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
        body.base_style = self._document.styles["Normal"]
        body.font.size = Pt(12)
        body.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        body.paragraph_format.first_line_indent = Mm(12.5)
        body.paragraph_format.space_after = Pt(0)
        body.paragraph_format.widow_control = True
        msn_message = self._document.styles.add_style("MSN Message", WD_STYLE_TYPE.PARAGRAPH)
        msn_message.font.name = "Arial"
        msn_message.font.size = Pt(10)
        msn_message.base_style = self._document.styles["Body"]
        msn_message.paragraph_format.first_line_indent = Mm(0)
        msn_sender = self._document.styles.add_style("MSN Sender", WD_STYLE_TYPE.PARAGRAPH)
        msn_sender.font.name = "Arial"
        msn_sender.font.size = Pt(10)
        msn_sender.base_style = self._document.styles["Body"]
        msn_sender.font.color.rgb = RGBColor.from_string(self._grey)
        msn_sender.paragraph_format.first_line_indent = Mm(0)
        msn_sender.paragraph_format.keep_with_next = True
        msn_sender.next_paragraph_style = msn_message

    def _process_raw_document_content(self):
        author = self._document.add_paragraph(self._raw_metadata["author"])
        author.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title = self._document.add_paragraph(self._raw_metadata["title"], self._document.styles["Title"])
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for line in self._raw_document_content:
            if line.startswith("# "):
                line = line.replace("# ", "")
                self._document.add_heading(line, level=1)
            elif line.startswith("András:") or line.startswith("Erika:"):
                name, line = line.split(": ")
                name += " üzenete:"
                self._document.add_paragraph(name, style=self._document.styles["MSN Sender"])
                current_p = self._document.add_paragraph("", style=self._document.styles["MSN Message"])
                bullet = current_p.add_run("\u2022\u2000")
                bullet.font.color.rgb = RGBColor.from_string(self._grey)
                current_p.add_run(line)
            else:
                self._document.add_paragraph(line, style=self._document.styles["Body"])
        year_p = self._document.add_paragraph()
        year_p.paragraph_format.space_before = Pt(11)
        year_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        year = year_p.add_run(self._raw_metadata["rights"][-4:])
        year.font.italic = True

if __name__ == "__main__":
    d = DocxConverter(argv[1])
    d.convert()